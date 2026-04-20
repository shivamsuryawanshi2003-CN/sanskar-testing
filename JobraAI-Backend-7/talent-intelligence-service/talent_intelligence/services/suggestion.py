import json
from typing import Any

try:
    from openai import OpenAI
except ImportError:  # pragma: no cover - depends on environment
    OpenAI = None

from talent_intelligence.services.heuristics import build_upgrade_heuristically
from talent_intelligence.services.llm_shared.common import (
    OPENAI_API_KEY,
    OPENAI_BASE_URL,
    OPENAI_MODEL,
    OPENAI_TIMEOUT_SECONDS,
    _extract_response_text,
    _responses_create_with_retry,
    openai_key_available,
)

def _validate_choice_id(choice_id: int) -> None:
    if choice_id not in (1, 2, 3):
        raise ValueError("choice_id must be 1, 2, or 3")


def _validate_upgrade_json(raw: Any) -> dict[str, Any]:
    if not isinstance(raw, dict):
        raise ValueError("Upgrade response must be a JSON object.")
    required_fields = [
        "Professional_Summary", "Work_Experience", "Skills", "Education", "Certifications",
        "ATS_Score", "Summary", "Suggestions_for_Improvement", "Score_Breakdown",
        "Missing_Keywords", "Achievements_or_Certifications", "Resume_Strength",
        "Key_Skills", "Overused_Keywords_Count", "Repeated_Word_Frequency", "Word_Replacement_Suggestions"
    ]
    for field in required_fields:
        if field not in raw:
            raise ValueError(f"Upgrade response missing {field}.")
    if not isinstance(raw["Professional_Summary"], str):
        raise ValueError("Professional_Summary must be a string.")
    if not isinstance(raw["Work_Experience"], list):
        raise ValueError("Work_Experience must be a list.")
    if not isinstance(raw["Skills"], dict):
        raise ValueError("Skills must be an object.")
    if not isinstance(raw["ATS_Score"], (int, float)):
        raise ValueError("ATS_Score must be a number.")
    return raw


def upgrade_resume_with_gpt35(
    resume_text: str,
    ats_report: dict[str, Any],
    choice_id: int,
    context_data: str | None = None,
) -> dict[str, Any]:
    """Generate resume upgrade JSON using OpenAI when available, else a deterministic fallback."""
    _validate_choice_id(choice_id)

    if not openai_key_available():
        return build_upgrade_heuristically(
            resume_text=resume_text,
            ats_report=ats_report,
            choice_id=choice_id,
            context_data=context_data or "",
        )
    if OpenAI is None:
        raise ValueError("openai package is not installed.")

    context_data = context_data or ""
    strategies = {
        1: ("GENERAL_REPAIR", "Fix structural health and action-verb density."),
        2: ("ROLE_ALIGNMENT", f"Optimize for industry standards for the role: {context_data}"),
        3: ("JD_MATCHING", f"Surgically inject keywords and align bullets with this JD: {context_data}"),
    }
    strategy_name, objective = strategies[choice_id]

    system_message = (
        "You are an automated JSON-only resume optimizer. "
        "You must output ONLY valid JSON. No conversational text. No markdown blocks. "
        "Every key must be in Capitalized_Snake_Case. "
        "If the input is not a resume, return {\"Error\": \"Invalid Document\"}."
    )

    user_content = f"""
### TASK: {strategy_name}
OBJECTIVE: {objective}

### DATA_INPUT
RESUME_TEXT: {resume_text[:8000]}
PHASE_1_ATS_REPORT: {json.dumps(ats_report, ensure_ascii=False)}

### CONSTRAINTS
- Do not invent new jobs.
- Rephrase existing bullets to fix weaknesses in the ATS report.
- Use Capitalized_Snake_Case for all JSON keys.
- Group skills by category.
- Include the current ATS score and all analysis fields in the upgraded resume output.

### OUTPUT_JSON_SCHEMA_EXAMPLE
{{
  "Professional_Summary": "...",
  "Work_Experience": [{{"Job_Title": "...", "Company": "...", "Dates": "...", "Description": "..."}}],
  "Skills": {{"Technical_Skills": [], "Soft_Skills": []}},
  "Education": ["..."],
  "Certifications": ["..."],
  "ATS_Score": 85,
  "Summary": ["Bullet point 1", "Bullet point 2"],
  "Suggestions_for_Improvement": ["Suggestion 1", "Suggestion 2"],
  "Score_Breakdown": {{"FORMAT": "90%", "SKILLS": "85%", "EXPERIENCE": "80%", "COMPLETENESS": "75%"}},
  "Missing_Keywords": {{"SUMMARY": 2, "SKILLS": 3, "EXPERIENCE": 1, "PROJECTS": 0}},
  "Achievements_or_Certifications": ["Certification 1", "Award 1"],
  "Resume_Strength": ["Strength 1", "Strength 2", "Strength 3"],
  "Key_Skills": ["Skill 1", "Skill 2", "Skill 3"],
  "Overused_Keywords_Count": 3,
  "Repeated_Word_Frequency": {{"managed": 6, "developed": 4}},
  "Word_Replacement_Suggestions": ["3 times: managed\\ntry replacing with\\nled\\ndirected\\noversaw", "4 times: developed\\ntry replacing with\\nbuilt\\ncreated\\ndesigned"]
}}
"""

    client = OpenAI(
        api_key=OPENAI_API_KEY,
        base_url=OPENAI_BASE_URL,
        timeout=OPENAI_TIMEOUT_SECONDS,
    )
    response = _responses_create_with_retry(
        client=client,
        model=OPENAI_MODEL,
        input_payload=[
            {"role": "system", "content": [{"type": "input_text", "text": system_message}]},
            {"role": "user", "content": [{"type": "input_text", "text": user_content}]},
        ],
        max_output_tokens=1200,
        store=False,
    )

    raw_text = _extract_response_text(response)
    if not raw_text:
        raise ValueError("OpenAI returned empty upgrade payload.")

    if raw_text.startswith("```"):
        raw_text = raw_text.split("```", 1)[1].rsplit("```", 1)[0].strip()
    elif "```json" in raw_text:
        raw_text = raw_text.split("```json", 1)[1].split("```", 1)[0].strip()

    try:
        upgraded_json = json.loads(raw_text)
    except json.JSONDecodeError as exc:
        raise ValueError(f"AI_RESPONSE_PARSING_FAILED: {str(exc)}") from exc

    return _validate_upgrade_json(upgraded_json)

def save_as_docx(json_data, filename):
    """
    Standard JOBRa AI DOCX Generation logic.
    """
    try:
        from docx import Document
    except ImportError:
        raise ValueError("python-docx is required for DOCX generation. Install with: pip install python-docx")
    
    if "Error" in json_data:
        return None
        
    doc = Document()
    for section, content in json_data.items():
        doc.add_heading(section.replace("_", " "), level=1)
        
        if isinstance(content, list):
            for item in content:
                if isinstance(item, dict):
                    for k, v in item.items():
                        p = doc.add_paragraph()
                        p.add_run(f"{k.replace('_', ' ')}: ").bold = True
                        p.add_run(str(v))
                else:
                    doc.add_paragraph(f"• {item}")
        elif isinstance(content, dict):
            for k, v in content.items():
                p = doc.add_paragraph()
                p.add_run(f"{k.replace('_', ' ')}: ").bold = True
                p.add_run(str(v))
        else:
            doc.add_paragraph(str(content))
            
    doc_path = f"Improved_{filename}.docx"
    doc.save(doc_path)
    return doc_path
