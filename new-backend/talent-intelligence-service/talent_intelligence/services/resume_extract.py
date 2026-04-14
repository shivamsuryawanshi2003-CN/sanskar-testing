from io import BytesIO


class ResumeExtractError(Exception):
    pass


ALLOWED_EXTENSIONS = {".pdf", ".docx"}


def extract_resume_text(filename: str, data: bytes) -> str:
    if not data:
        raise ResumeExtractError("Empty file upload.")
    lower = filename.lower().strip()
    if lower.endswith(".pdf"):
        return _extract_pdf(data)
    if lower.endswith(".docx"):
        return _extract_docx(data)
    raise ResumeExtractError("Only .pdf and .docx resume files are accepted.")


def _extract_pdf(data: bytes) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise ResumeExtractError("pdfplumber is required for PDF parsing.") from exc
    try:
        parts: list[str] = []
        with pdfplumber.open(BytesIO(data)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(text.strip())
        out = "\n".join(parts).strip()
        if not out:
            raise ResumeExtractError("No extractable text found in PDF.")
        return out
    except ResumeExtractError:
        raise
    except Exception as exc:
        raise ResumeExtractError("Could not parse PDF file.") from exc


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ResumeExtractError("python-docx is required for DOCX parsing.") from exc
    try:
        doc = Document(BytesIO(data))
        parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        out = "\n".join(parts).strip()
        if not out:
            raise ResumeExtractError("No extractable text found in DOCX.")
        return out
    except ResumeExtractError:
        raise
    except Exception as exc:
        raise ResumeExtractError("Could not parse DOCX file.") from exc
